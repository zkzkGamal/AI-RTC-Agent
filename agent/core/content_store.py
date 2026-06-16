"""
Content Store (content_store.py)
--------------------------------
Single source of truth for the global ``content/`` folder that lives at the
project root. This is where uploaded CVs (and any other reference documents)
are saved, and where the agent looks for a CV when the user only mentions it
by name in the CLI / chat.

Responsibilities:
- Expose the absolute path to the global ``content/`` directory (created on import).
- Resolve a user-provided file name / partial path to a real file inside ``content/``.
- Read text from supported document types ONLY: PDF, Word (.docx), and Markdown/text.
"""
import pathlib
import logging

logger = logging.getLogger(__name__)

# project_root/content  (agent/core/content_store.py -> parents[2] == project root)
PROJECT_ROOT = pathlib.Path(__file__).resolve().parents[2]
CONTENT_DIR = PROJECT_ROOT / "content"
CONTENT_DIR.mkdir(parents=True, exist_ok=True)

# Only these document types may be read.
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


def is_supported(file_path: str) -> bool:
    return pathlib.Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS


def save_bytes(filename: str, data: bytes) -> pathlib.Path:
    """Persist raw bytes into the content folder, returning the saved path."""
    safe_name = pathlib.Path(filename).name  # strip any directory components
    if not is_supported(safe_name):
        raise ValueError(
            f"Unsupported file type '{pathlib.Path(safe_name).suffix}'. "
            f"Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    dest = CONTENT_DIR / safe_name
    dest.write_bytes(data)
    logger.info(f"[content_store] Saved {len(data)} bytes to {dest}")
    return dest


def resolve_path(file_path: str | None) -> pathlib.Path | None:
    """
    Resolve a user-supplied reference to an actual file.

    - ``None`` / empty                -> most recently modified supported file in content/.
    - an existing absolute/relative path -> used as-is.
    - a bare name (e.g. "cv.pdf" or "cv") -> matched inside content/.
    """
    if not file_path or not str(file_path).strip():
        return latest_document()

    raw = str(file_path).strip().strip('"').strip("'")
    candidate = pathlib.Path(raw)

    # 1. Direct hit (absolute or relative to CWD).
    if candidate.is_file():
        return candidate

    # 2. Look inside the content folder by exact name.
    by_name = CONTENT_DIR / candidate.name
    if by_name.is_file():
        return by_name

    # 3. Fuzzy match: stem appears in a content file name (handles "cv" -> "my_cv.pdf").
    stem = candidate.stem.lower()
    matches = [
        p for p in CONTENT_DIR.iterdir()
        if p.is_file() and is_supported(p.name) and stem in p.stem.lower()
    ]
    if matches:
        return max(matches, key=lambda p: p.stat().st_mtime)

    return None


def latest_document() -> pathlib.Path | None:
    """Return the most recently modified supported document in the content folder."""
    docs = [p for p in CONTENT_DIR.iterdir() if p.is_file() and is_supported(p.name)]
    if not docs:
        return None
    return max(docs, key=lambda p: p.stat().st_mtime)


def read_document(file_path: str | None) -> str:
    """
    Read and return the plain text of a supported document (PDF / Word / Markdown).

    Raises FileNotFoundError if nothing matches and ValueError for unsupported types.
    """
    resolved = resolve_path(file_path)
    if resolved is None:
        raise FileNotFoundError(
            f"No matching document found for '{file_path}'. "
            f"Place the file in the content/ folder ({CONTENT_DIR})."
        )
    if not is_supported(resolved.name):
        raise ValueError(
            f"Unsupported file type '{resolved.suffix}'. "
            f"Only PDF, Word (.docx) and Markdown/text files can be read."
        )

    suffix = resolved.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(resolved)
    elif suffix == ".docx":
        text = _read_docx(resolved)
    else:  # .md / .markdown / .txt
        text = resolved.read_text(encoding="utf-8", errors="ignore")

    text = text.strip()
    if not text:
        raise ValueError(f"Document '{resolved.name}' contained no extractable text.")
    return text


def _read_pdf(path: pathlib.Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: pathlib.Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ImportError(
            "Reading .docx files requires the 'python-docx' package. "
            "Install it with: pip install python-docx"
        ) from e

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Include table cell text (common in CV templates).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
